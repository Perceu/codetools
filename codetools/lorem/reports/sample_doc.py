import docx
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from io import BytesIO
from django.conf import settings


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def sampleDoc():

    document = Document()
    document.add_heading('Code Tools - Sample DOC', 0)
    document.add_heading('Code Tools v {}'.format(settings.APP_VERSION), level=3)
    document.add_paragraph('Documento criado como sample de DOC. ')
    document.add_paragraph('Para mais informações visite. ')
    document.add_paragraph('codetools.com.br')
    save_document = BytesIO()
    document.save(save_document)
    return save_document